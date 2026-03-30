import os
import re
import json
import logging
import argparse
from datetime import datetime
from typing import List, Dict, Any

import docx
import numpy as np
from sentence_transformers import SentenceTransformer

# Import existing FAISS wrapper
import sys
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
sys.path.append(os.path.join(PROJECT_ROOT, "rag_api"))
from utils.vector_db import get_vector_db

# Configuration
METADATA_JSON_PATH = os.path.join(PROJECT_ROOT, "rag_api", "faiss_index", "itinerary_metadata.json")
SUMMARY_LOG_PATH = os.path.join(PROJECT_ROOT, "data", "processed", "itinerary_ingest_summary.json")
STOP_SECTION_KEYWORDS = ["important information", "visa information", "general information", "terms and conditions"]

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

class ItineraryIngestor:
    def __init__(self):
        self.db = get_vector_db()
        self.processed_files = self._load_existing_metadata()

    def _load_existing_metadata(self) -> set:
        if os.path.exists(METADATA_JSON_PATH):
            try:
                with open(METADATA_JSON_PATH, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    return {item["filename"] for item in data}
            except Exception as e:
                logger.error(f"Failed to load existing metadata: {e}")
        return set()

    def _save_metadata(self, new_items: List[Dict[str, Any]]):
        all_metadata = []
        if os.path.exists(METADATA_JSON_PATH):
            with open(METADATA_JSON_PATH, "r", encoding="utf-8") as f:
                all_metadata = json.load(f)
        
        all_metadata.extend(new_items)
        os.makedirs(os.path.dirname(METADATA_JSON_PATH), exist_ok=True)
        with open(METADATA_JSON_PATH, "w", encoding="utf-8") as f:
            json.dump(all_metadata, f, indent=2)

    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        doc = docx.Document(file_path)
        filename = os.path.basename(file_path)
        
        full_text_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # 1. Extract Overview Details
        tour_name = full_text_paras[0] if full_text_paras else "Unknown Tour"
        destinations = []
        duration = ""
        
        for para in full_text_paras:
            if "Destinations Covered:" in para:
                dests_str = para.split("Destinations Covered:")[1].strip()
                destinations = [d.strip() for d in re.split(r"[-|,]", dests_str) if d.strip()]
            if "Duration:" in para:
                duration = para.split("Duration:")[1].strip()

        # 2. Extract Inclusions/Exclusions
        inclusions = []
        exclusions = []
        in_incl = False
        in_excl = False
        
        for para in full_text_paras:
            para_upper = para.upper()
            if "PRICE INCLUDES" in para_upper or "INCLUSIONS" in para_upper:
                in_incl = True
                in_excl = False
                continue
            if "NOT INCLUDE" in para_upper or "EXCLUSIONS" in para_upper:
                in_incl = False
                in_excl = True
                continue
            if any(kw.upper() in para_upper for kw in STOP_SECTION_KEYWORDS):
                break
                
            if in_incl and para: inclusions.append(para)
            if in_excl and para: exclusions.append(para)

        # 3. Extract Day Chunks
        days = []
        current_day = None
        
        for para in full_text_paras:
            if any(kw.upper() in para.upper() for kw in STOP_SECTION_KEYWORDS):
                break
                
            day_match = re.match(r"Day\s*(\d+)", para, re.IGNORECASE)
            if day_match:
                if current_day:
                    days.append(current_day)
                current_day = {"day": day_match.group(1), "text": para}
            elif current_day:
                current_day["text"] += " " + para
        
        if current_day:
            days.append(current_day)

        # 4. Extract Hotels Table
        hotels_prose = []
        for table in doc.tables:
            # Look for a table with "Hotel" or "City" in the first row
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if "hotel" in headers or "city" in headers or "nights" in headers:
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 3:
                        # Assuming City, Hotel, Nights order or similar
                        city, hotel, nights = cells[0], cells[1], cells[2]
                        if city and hotel:
                            hotels_prose.append(f"In {city}, guests stay at {hotel} for {nights} nights.")
                break # Only process the first relevant table

        return {
            "filename": filename,
            "tour_name": tour_name,
            "destinations": destinations,
            "duration": duration,
            "days": days,
            "overview_parts": {
                "inclusions": inclusions,
                "exclusions": exclusions
            },
            "hotels": " ".join(hotels_prose)
        }

    def chunk_text(self, text: str, max_words: int = 150) -> List[str]:
        words = text.split()
        if len(words) <= max_words:
            return [text]
        
        chunks = []
        # Split by sentences (simple period followed by space)
        sentences = re.split(r"(?<=[.!?])\s+", text)
        current_chunk = []
        current_len = 0
        
        for sentence in sentences:
            sent_len = len(sentence.split())
            if current_len + sent_len > max_words and current_chunk:
                chunks.append(" ".join(current_chunk))
                current_chunk = [sentence]
                current_len = sent_len
            else:
                current_chunk.append(sentence)
                current_len += sent_len
        
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        return chunks

    def ingest_folder(self, input_dir: str):
        if not os.path.exists(input_dir):
            logger.error(f"Input directory does not exist: {input_dir}")
            return

        docx_files = [f for f in os.listdir(input_dir) if f.endswith(".docx") and not f.startswith("~$")]
        logger.info(f"Found {len(docx_files)} .docx files in {input_dir}")

        new_metadata_entries = []
        summary = {
            "files_processed": [],
            "chunks_per_type": {"day": 0, "overview": 0, "hotels": 0},
            "total_vectors_added": 0,
            "failed_files": []
        }

        for filename in docx_files:
            if filename in self.processed_files:
                logger.info(f"Skipping already ingested file: {filename}")
                continue

            file_path = os.path.join(input_dir, filename)
            try:
                parsed_data = self.parse_docx(file_path)
                
                texts_to_add = []
                metadatas_to_add = []
                
                # a) Day Chunks
                for day_data in parsed_data["days"]:
                    day_text = day_data["text"]
                    # Preserve destination/context
                    context_prefix = f"{parsed_data['tour_name']} - "
                    
                    for chunk in self.chunk_text(day_text):
                        final_text = context_prefix + chunk
                        texts_to_add.append(final_text)
                        metadatas_to_add.append({
                            "source": "itinerary",
                            "filename": filename,
                            "tour_name": parsed_data["tour_name"],
                            "destinations": parsed_data["destinations"],
                            "duration": parsed_data["duration"],
                            "chunk_type": "day"
                        })
                        summary["chunks_per_type"]["day"] += 1

                # b) Overview Chunk
                overview_text = (
                    f"Tour: {parsed_data['tour_name']}. "
                    f"Destinations: {', '.join(parsed_data['destinations'])}. "
                    f"Duration: {parsed_data['duration']}. "
                    f"Inclusions: {' '.join(parsed_data['overview_parts']['inclusions'][:15])}. "
                    f"Exclusions: {' '.join(parsed_data['overview_parts']['exclusions'][:15])}."
                )
                texts_to_add.append(overview_text)
                metadatas_to_add.append({
                    "source": "itinerary",
                    "filename": filename,
                    "tour_name": parsed_data["tour_name"],
                    "destinations": parsed_data["destinations"],
                    "duration": parsed_data["duration"],
                    "chunk_type": "overview"
                })
                summary["chunks_per_type"]["overview"] += 1

                # c) Hotels Chunk
                if parsed_data["hotels"]:
                    texts_to_add.append(f"{parsed_data['tour_name']} Hotels: " + parsed_data["hotels"])
                    metadatas_to_add.append({
                        "source": "itinerary",
                        "filename": filename,
                        "tour_name": parsed_data["tour_name"],
                        "destinations": parsed_data["destinations"],
                        "duration": parsed_data["duration"],
                        "chunk_type": "hotels"
                    })
                    summary["chunks_per_type"]["hotels"] += 1

                # Add to FAISS
                if texts_to_add:
                    self.db.add_texts(texts_to_add, metadatas_to_add)
                    summary["total_vectors_added"] += len(texts_to_add)
                
                new_metadata_entries.append({
                    "filename": filename,
                    "ingested_at": datetime.now().isoformat(),
                    "tour_name": parsed_data["tour_name"]
                })
                summary["files_processed"].append(filename)
                logger.info(f"Successfully ingested: {filename}")

            except Exception as e:
                logger.error(f"Failed to process {filename}: {e}")
                summary["failed_files"].append({"filename": filename, "reason": str(e)})

        # Save updates
        if new_metadata_entries:
            self._save_metadata(new_metadata_entries)
            
        # Log summary
        os.makedirs(os.path.dirname(SUMMARY_LOG_PATH), exist_ok=True)
        with open(SUMMARY_LOG_PATH, "w", encoding="utf-8") as f:
            json.dump(summary, f, indent=2)
            
        logger.info(f"Ingestion complete. Added {summary['total_vectors_added']} vectors.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Ingest Word itineraries into FAISS index.")
    parser.add_argument("--input_dir", required=True, help="Directory containing .docx files")
    args = parser.parse_args()

    # Ensure project root is in path for module imports
    sys.path.append(PROJECT_ROOT)
    
    ingestor = ItineraryIngestor()
    ingestor.ingest_folder(args.input_dir)
